import flet as ft
import shutil
import os
from classes.braille import BrailleConverter
import string
import random
from docx import Document
from docx.shared import Inches


class Braille(ft.UserControl):
    def __init__(self, page: ft.Page):
        super().__init__()
        self.page = page
        self.letters = string.ascii_letters
        self.letter_blocks = ft.Text()
        self.red = random.randint(0, 255)
        self.green = random.randint(0, 255)
        self.blue = random.randint(0, 255)
        self.color = self.random_color()
        self.braille = ft.Text()
        self.braille_words = ft.TextField(
            width=700,
            multiline=True,
            max_lines=10,
            border_color="#007BDD",
            keyboard_type=ft.KeyboardType.NAME,
            border_radius=ft.border_radius.all(10),
            hint_text="enter some text".capitalize(),
            prefix_icon=ft.icons.DIALPAD_ROUNDED,
        )
        self.braille_converter = BrailleConverter()

        #  ---------------------------the grid view that will keep the container for the brail code---//
        self.dictionary_words = ft.GridView(
            expand=1,
            runs_count=6,
            max_extent=70,
            child_aspect_ratio=0.9,
            spacing=1,
            run_spacing=5,
            padding=30,
            width=300,
        )
        #   --------------------the generated braille codes here--------------//
        self.file_picker = ft.FilePicker(on_result=self.upload_files)
        self.page.overlay.append(self.file_picker)
        self.document_name = ft.TextField(
            width=500,
            border_color="#0050C1",
            label="enter the document_name",
            hint_text="document name",
            autofocus=True,
            helper_text="document name",
            prefix_icon=ft.icons.SCATTER_PLOT_ROUNDED,
            focused_border_color="#1a237e", focused_color="#6200ea",
        )
        self.dictionary_dialog = ft.AlertDialog()

    #   -----------------validating the input field here-----------------//
    def validate_input(self, e):
        try:
            if not self.braille_words.value:
                self.braille_words.error_text = "fill in something".capitalize()
                self.update()
            else:
                self.show_generated_text()
        except Exception as ex:
            print(ex)

    # Generate a random RGB color
    def random_color(self):
        return self.red, self.green, self.blue

    #   ----------------------function to upload files here-----------//
    def upload_files(self, e: ft.FilePickerResultEvent):
        try:
            for x in e.files:
                print(x.path)
                print(x.name)
                # THIS FOR GET CURRENT LOCATION NOW
                you_copy = os.path.join(os.getcwd(), "assets")
                shutil.copy(x.path, you_copy)
                self.update()
        except Exception as ex:
            print(ex)

    #  ----------------show dialog text here----------------//
    def show_generated_text(self):
        try:
            self.braille = self.braille_converter.generate_braille(self.braille_words.value)
            self.dictionary_dialog = ft.AlertDialog(
                content=ft.Container(
                    border_radius=ft.border_radius.all(10),
                    width=900,
                    bgcolor="white",
                    padding=ft.padding.only(left=20, top=20),
                    expand=True,
                    content=ft.Column(
                        scroll=ft.ScrollMode.HIDDEN,
                        controls=[
                            ft.Container(
                                padding=ft.padding.only(left=10, right=10),
                                margin=ft.margin.only(top=15),
                                content=ft.Text(
                                    "generated braille".capitalize(),
                                    style=ft.TextThemeStyle.DISPLAY_SMALL,
                                    weight=ft.FontWeight.W_700,
                                    color="white"
                                ),
                            ),

                            ft.Text(
                                f"{self.braille}",
                                style=ft.TextThemeStyle.BODY_LARGE,
                                color="#3e2723",
                                size=25,
                                weight=ft.FontWeight.W_700
                            ),
                            ft.Container(
                                margin=ft.margin.only(top=20),
                                content=ft.Row(
                                    controls=[
                                        self.document_name
                                    ]
                                )
                            ),
                            ft.Container(
                                margin=ft.margin.only(bottom=20),
                                content=ft.Row(
                                    controls=[
                                        ft.ElevatedButton(
                                            on_click=self.export_to_word,
                                            bgcolor="#007BDD",
                                            content=ft.Row(
                                                controls=[
                                                    ft.Icon(
                                                        ft.icons.BROADCAST_ON_HOME_ROUNDED,
                                                        size=20,
                                                        color="white"),
                                                    ft.Text(
                                                        "save to word".title(),
                                                        color="white"
                                                    )
                                                ],
                                            ),
                                        ),
                                    ]
                                )
                            )
                        ]
                    )
                )
            )
            self.page.dialog = self.dictionary_dialog
            self.dictionary_dialog.open = True
            self.page.update()
        except Exception as ex:
            print(ex)

    def on_click(self, e):
        self.letter_blocks = e.control.data
        braille = self.braille_converter.generate_braille(self.letter_blocks)
        dictionary_dialog = ft.AlertDialog(
            content=ft.Container(
                width=500,
                height=300,
                content=ft.Column(
                    scroll=ft.ScrollMode.HIDDEN,
                    controls=[
                        ft.Row(
                            vertical_alignment=ft.CrossAxisAlignment.CENTER,
                            alignment=ft.MainAxisAlignment.CENTER,
                            controls=[
                                ft.Container(
                                    margin=ft.margin.only(top=100),
                                    content=ft.Column(
                                        controls=[
                                            ft.Text(
                                                f"{braille}",
                                                style=ft.TextThemeStyle.BODY_LARGE,
                                                color=f"{self.color}",
                                                size=90,
                                                weight=ft.FontWeight.W_700
                                            ),
                                            #  -----------------the container that converts the braille to word document---/
                                        ]
                                    )
                                )
                            ]
                        )
                    ]
                )
            )
        )
        self.page.dialog = dictionary_dialog
        dictionary_dialog.open = True
        self.page.update()

    #   --------------------------the keyboard start here for the system----------------//
    def alphabet_letters(self):
        try:
            for letter in self.letters:
                self.dictionary_words.controls.append(
                    ft.Card(
                        elevation=1.0,
                        content=ft.Container(
                            data=letter,
                            content=ft.Column(
                                controls=[
                                    ft.Row(
                                        alignment=ft.MainAxisAlignment.CENTER,
                                        controls=[
                                            ft.Text(f"{letter}", size=30, weight=ft.FontWeight.W_500, color="#212121")]
                                    )
                                ]
                            ),
                            on_click=self.on_click
                        )
                    )
                )

        except Exception as ex:
            print(ex)

    def build(self):
        self.alphabet_letters()
        return ft.ListView(
            expand=1,
            auto_scroll=True,
            spacing=10,
            width=1200,
            height=700,
            scale=1.0,
            controls=[
                ft.Container(
                    padding=ft.padding.only(left=10, right=10),
                    margin=ft.margin.only(top=15),
                    content=ft.Column(
                        controls=[
                            # -----------------the top page for the controls here-----//
                            ft.Container(
                                margin=ft.margin.only(top=30, left=20, right=20),
                                content=ft.Row(
                                    alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                                    controls=[
                                        ft.Text(
                                            "AI braille copilot".title(),
                                            style=ft.TextThemeStyle.DISPLAY_SMALL,
                                            weight=ft.FontWeight.W_500,
                                            color="#EF3A71"
                                        ),
                                    ]
                                )
                            ),

                            #  ---------------------the welcome text containers here-------------------//
                            ft.Container(
                                margin=ft.margin.only(top=30),
                                content=ft.Row(
                                    alignment=ft.MainAxisAlignment.CENTER,
                                    controls=[
                                        ft.Card(
                                            content=ft.Container(
                                                bgcolor="#212121",
                                                border_radius=ft.border_radius.all(10),
                                                content=ft.Column(
                                                    controls=[
                                                        ft.Container(
                                                            margin=ft.margin.only(top=20),
                                                            padding=ft.padding.only(left=200),
                                                            content=ft.Row(
                                                                alignment=ft.MainAxisAlignment.CENTER,
                                                                controls=[
                                                                    ft.Icon(
                                                                        ft.icons.LIGHT_MODE_ROUNDED,
                                                                        size=40,
                                                                        color="#FFBB4D"
                                                                    )
                                                                ]
                                                            )
                                                        ),
                                                        ft.Container(
                                                            padding=ft.padding.only(left=10, right=10),
                                                            margin=ft.margin.only(bottom=20),
                                                            width=450,
                                                            content=ft.Column(
                                                                controls=[
                                                                    ft.Text(
                                                                        "the system will translate your uploaded"
                                                                        "text inputs to braille codes and later printed."
                                                                        "the braille dots will be generated with the"
                                                                        "accuracy of 98.9% as predicted by the machine "
                                                                        "learning models".capitalize(),
                                                                        weight=ft.FontWeight.W_500,
                                                                        size=15,
                                                                        color="white"
                                                                    )
                                                                ]
                                                            )
                                                        )
                                                    ]
                                                )
                                            )
                                        ),
                                        #  -------------------the other container car will be here------------//
                                        ft.Card(
                                            content=ft.Container(
                                                border_radius=ft.border_radius.all(10),
                                                gradient=ft.LinearGradient(
                                                    colors=[
                                                        "#00B4C6",
                                                        "#007BDD"
                                                    ],
                                                    begin=ft.alignment.bottom_left,
                                                    end=ft.alignment.top_right
                                                ),
                                                content=ft.Column(
                                                    controls=[
                                                        ft.Container(
                                                            margin=ft.margin.only(top=20),
                                                            padding=ft.padding.only(left=200),
                                                            content=ft.Row(
                                                                alignment=ft.MainAxisAlignment.CENTER,
                                                                controls=[
                                                                    ft.Icon(
                                                                        ft.icons.MULTITRACK_AUDIO_ROUNDED,
                                                                        size=40,
                                                                        color="#FFBB4D"
                                                                    )
                                                                ]
                                                            )
                                                        ),
                                                        ft.Container(
                                                            padding=ft.padding.only(left=10, right=10),
                                                            margin=ft.margin.only(bottom=20),
                                                            width=450,
                                                            content=ft.Column(
                                                                controls=[
                                                                    ft.Text(
                                                                        "all the corrected translations will be stored"
                                                                        "for future referencing and also to improve"
                                                                        "how the system functions in the future to "
                                                                        "improve the readability, and accuracy which is "
                                                                        "the core part of the system".capitalize(),
                                                                        weight=ft.FontWeight.W_500,
                                                                        size=15,
                                                                        color="white"
                                                                    )
                                                                ]
                                                            )
                                                        )
                                                    ]
                                                )
                                            )
                                        ),
                                    ]
                                )
                            ),
                            #   ----------------the container for the keyboard will be here----------------//
                            ft.Container(
                                content=ft.Column(
                                    controls=[
                                        ft.Card(
                                            content=ft.Container(
                                                content=ft.Row(
                                                    controls=[
                                                        self.dictionary_words
                                                    ]
                                                ),
                                            ),
                                        ),
                                        #  ----------------the container to add the text boxes--------------//
                                        ft.Container(
                                            padding=ft.padding.only(left=10, right=10),
                                            margin=ft.margin.only(top=15),
                                            content=ft.Column(
                                                controls=[
                                                    ft.Row(
                                                        controls=[
                                                            self.braille_words,
                                                            ft.IconButton(
                                                                bgcolor="#007BDD",
                                                                height=50,
                                                                width=50,
                                                                on_click=self.validate_input,
                                                                icon=ft.icons.SEND_ROUNDED,
                                                                icon_color="white",
                                                                tooltip="translate text"
                                                            ),
                                                        ]
                                                    ),

                                                ]
                                            )
                                        )
                                        #  ---------------------end of the container here---------------//
                                    ]
                                )
                            )

                        ]
                    )
                )
            ]
        )

    def export_to_word(self, e):
        try:

            document = Document()
            # add a heading to the document
            document.add_heading('SOCHIE TECHNICAL COLLEGE', 0)
            # add a paragraph to the document
            p1 = document.add_paragraph('THE LECTURE LESSON')
            # add another paragraph to the document
            p2 = document.add_paragraph('This is another paragraph.')

            p3 = document.add_paragraph(f"{self.braille}")
            # add an image to the document
            img_path = 'assets/signs/pexels-andre-moura-2402467.jpg'
            document.add_picture(img_path, width=Inches(2))
            # save the document
            document.save(f'C:/FinalProject/Lessons/{self.document_name.value}.docx')
            self.dictionary_dialog.open = False
            self.page.update()
            self.page.snack_bar = ft.SnackBar(
                bgcolor="#0050C1",
                content=ft.Container(
                    content=ft.Row(
                        controls=[
                            ft.Text(
                                "your document has been saved successfully".capitalize(),
                                style=ft.TextThemeStyle.DISPLAY_SMALL,
                                weight=ft.FontWeight.W_700,
                                color="white"
                            )
                        ]
                    )
                )
            )
            self.page.snack_bar.open = True
            self.page.update()
        except Exception as ex:
            print(ex)

