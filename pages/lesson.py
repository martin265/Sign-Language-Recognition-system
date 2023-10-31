import flet as ft
from classes.media_notes import AudioRecorder
import pyaudio
import speech_recognition as s_r
import random
from connection.db import my_connection
import openai
import wave
import os
import datetime
import sounddevice as sd
import soundfile as sf
from docx import Document
from docx.shared import Inches

# ------------------------the chart state here--------------------//

openai.api_key = open("pages/Api.txt", "r").read().strip("\n")


class Lessons(ft.UserControl):
    def __init__(self, page: ft.Page):
        super().__init__()
        self.page = page
        self.RECORDING_PATH = "C:/FinalProject/recordings"
        self.database_cursor = my_connection
        self.recording_name = ft.Text()
        self.file_name = "../output.wav"
        self.translated_text = ft.AlertDialog()
        #  ---------------------------the class for praying the audio files here---------//
        self.recording_file_name = ft.TextField(
            width=500,
            border_color="#0050C1",
            label="enter the recording name".title(),
            hint_text="recording name",
            autofocus=True,
            prefix_icon=ft.icons.MIC_ROUNDED,
            helper_text="enter the filename".title(),
        )
        self.audio1 = ft.Audio(
            src=self.file_name,
            autoplay=False,
            volume=1,
            balance=0,
        )
        self.page.overlay.append(self.audio1)
        self.microphone = ft.Dropdown(
            width=500,
            border_color="#0050C1",
            label="select microphone...",
            hint_text="active mic",
            autofocus=True,
            prefix_icon=ft.icons.MIC_ROUNDED,
            helper_text="select active microphone...",
        )
        self.channel = ft.Dropdown(
            width=500,
            border_color="#0050C1",
            label="select channel...",
            hint_text="active audio channel",
            autofocus=True,
            prefix_icon=ft.icons.WIFI_CHANNEL_ROUNDED,
            helper_text="select channel...",
        )
        self.recording_timer = ft.Dropdown(
            width=500,
            border_color="#0050C1",
            label="select microphone...",
            hint_text="active mic",
            autofocus=True,
            helper_text="select time(S)...",
            prefix_icon=ft.icons.TIMER_ROUNDED,
            focused_border_color="#1a237e", focused_color="#6200ea",
            options=[

            ]
        )
        self.current_lecture = ft.Dropdown(
            width=500,
            border_color="#0050C1",
            label="lecture",
            hint_text="select name",
            autofocus=True,
            helper_text="select name",
            prefix_icon=ft.icons.PERSON_2_ROUNDED,
            focused_border_color="#1a237e",
            focused_color="#6200ea",
            options=[

            ]
        )
        #  ---------------grid view for the recordings here---------------//
        self.all_recordings = ft.GridView(
            expand=1,
            runs_count=6,
            max_extent=300,
            child_aspect_ratio=0.9,
            spacing=1,
            run_spacing=5,
            padding=30,
            width=300,
        )
        self.duration = float(30)

        self.sound_recorder_dlg = ft.AlertDialog(
            content=ft.Container(
                content=ft.Text(
                    "recording started"
                )
            )
        )

        self.document_name = ft.TextField(
            width=500,
            border_color="#0050C1",
            label="enter the document_name",
            hint_text="document name",
            autofocus=True,
            helper_text="select time(S)...",
            prefix_icon=ft.icons.DOCUMENT_SCANNER_ROUNDED,
            focused_border_color="#1a237e", focused_color="#6200ea",
        )

    #   ----------------------the function to get active microphones-----------//
    def show_active_microphones(self):
        try:
            for single_microphone in s_r.Microphone.list_microphone_names():
                self.microphone.options.append(
                    ft.dropdown.Option(
                        single_microphone
                    )
                )
        except Exception as ex:
            print(ex)

    #   --------------------------show active channels-------------------//
    def show_active_channels(self):
        try:
            py_audios = pyaudio.PyAudio()
            for channel in py_audios.get_default_input_device_info():
                self.channel.options.append(
                    ft.dropdown.Option(
                        channel
                    )
                )
        except Exception as ex:
            print(ex)

    #   -------------------the timer for the recording here-------------//
    def show_recoding_timer(self):
        try:
            for timer in range(1, 30):
                self.recording_timer.options.append(
                    ft.dropdown.Option(
                        timer
                    )
                )
        except Exception as ex:
            print(ex)

    #   ---------------------------// show_-------------current user-----------//
    def show_logged_user(self):
        try:
            sql = "SELECT first_name FROM lecture"
            cursor = my_connection.cursor()
            cursor.execute(sql)
            all_results = cursor.fetchall()
            #  ----------pushing the data to the main table here----------------//
            columns = [column[0] for column in cursor.description]
            rows = [dict(zip(columns, row)) for row in all_results]
            #  ----------------looping through all the records here-------------//
            for lecture_name in rows:
                for keys in lecture_name.values():
                    self.current_lecture.options.append(
                        ft.dropdown.Option(
                            keys
                        )
                    )
        except Exception as ex:
            print(ex)

    def initializing_recording_session(self, e):
        """the method will be used to launch the recording session"""
        try:
            if not self.microphone.value:
                self.microphone.error_text = "choose a microphone first".capitalize()
                self.update()
            elif not self.channel.value:
                self.channel.error_text = "select a channel".capitalize()
                self.update()
            elif not self.recording_timer.value:
                self.recording_timer.error_text = "choose the duration".title()
                self.update()
            elif not self.current_lecture.value:
                self.current_lecture.error_text = "chose the lecture".title()
                self.update()
            else:
                self.start_recorder()
                self.fetch_all_recording()
        except Exception as ex:
            print(ex)

    def start_recorder(self):
        try:
            sample_rate = 44100
            channels = 1
            file_format = "wav"

            # Create the recordings directory if it does not exist
            if not os.path.exists("recordings"):
                os.makedirs("recordings")

            # Start recording
            self.page.snack_bar = ft.SnackBar(
                bgcolor="#0050C1",
                content=ft.Container(
                    content=ft.Row(
                        controls=[
                            ft.Text(
                                "your recording has started".capitalize(),
                                style=ft.TextThemeStyle.DISPLAY_SMALL,
                                weight=ft.FontWeight.W_700,
                                color="white"
                            )
                        ]
                    )
                )
            )
            self.page.snack_bar.open = True
            self.page.update()
            recording = sd.rec(int(self.duration * sample_rate), samplerate=sample_rate, channels=channels)
            # Wait for the recording to finish
            sd.wait()
            # Save the recording to a file
            file_name = os.path.join("recordings", f"{self.recording_file_name.value}.{file_format}")
            sf.write(file_name, recording, sample_rate)
            self.page.snack_bar = ft.SnackBar(
                bgcolor="#0050C1",
                content=ft.Container(
                    content=ft.Row(
                        controls=[
                            ft.Text(
                                f"your recording has completed and saved to {file_name}".capitalize(),
                                style=ft.TextThemeStyle.DISPLAY_SMALL,
                                weight=ft.FontWeight.W_700,
                                color="white"
                            )
                        ]
                    )
                )
            )
            self.fetch_all_recording()
            self.page.snack_bar.open = True
            self.page.update()
        except Exception as ex:
            print(ex)

    def volume_down(self, _):
        self.audio1.volume -= 0.1
        self.audio1.update()

    def volume_up(self, _):
        self.audio1.volume += 0.1
        self.audio1.update()

    def build(self):
        self.show_active_microphones()
        self.show_active_channels()
        self.show_recoding_timer()
        self.show_logged_user()
        self.fetch_all_recording()
        return ft.ListView(
            expand=True,
            auto_scroll=True,
            spacing=10,
            width=1200,
            height=700,
            scale=1.0,
            controls=[
                ft.Container(
                    content=ft.Column(
                        controls=[

                            ft.Container(
                                content=ft.Row(
                                    controls=[
                                        self.all_recordings
                                    ]
                                )
                            ),
                            ft.Container(
                                margin=ft.margin.only(left=30, right=30, top=30),
                                content=ft.Row(
                                    alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                                    controls=[
                                        ft.Text(
                                            "record lecture lesson".capitalize(),
                                            style=ft.TextThemeStyle.DISPLAY_SMALL,
                                            weight=ft.FontWeight.W_500,
                                            color="#0078D9"
                                        ),
                                        ft.ElevatedButton(
                                            on_click={},
                                            bgcolor="#0078D9",
                                            content=ft.Row(
                                                controls=[
                                                    ft.Icon(ft.icons.REFRESH_ROUNDED, size=20,
                                                            color="white"),
                                                    ft.Text(
                                                        "refresh details".title(),
                                                        color="white"
                                                    )
                                                ],

                                            ),
                                        )
                                    ]
                                )
                            ),
                            ft.Container(
                                border_radius=ft.border_radius.all(10),
                                padding=ft.padding.only(left=10, right=10, top=30, bottom=30),
                                margin=ft.margin.only(right=20),
                                bgcolor="#f5f5f5",
                                content=ft.Row(
                                    alignment=ft.MainAxisAlignment.CENTER,
                                    controls=[
                                        #  ----------------------for the form to record audio-----------//
                                        ft.Container(
                                            width=1000,
                                            content=ft.Column(
                                                controls=[
                                                    ft.Row([self.microphone, self.channel]),
                                                    ft.Row([self.recording_timer, self.current_lecture]),
                                                    ft.Row([self.recording_file_name]),
                                                    #   -------------------the row for the controls here------//
                                                    ft.Row(
                                                        controls=[
                                                            ft.ElevatedButton(
                                                                on_click=self.initializing_recording_session,
                                                                bgcolor="#007BDD",
                                                                content=ft.Row(
                                                                    controls=[
                                                                        ft.Icon(
                                                                            ft.icons.BROADCAST_ON_HOME_ROUNDED,
                                                                            size=20,
                                                                            color="white"),
                                                                        ft.Text(
                                                                            "start session".title(),
                                                                            color="white"
                                                                        )
                                                                    ],
                                                                ),
                                                            ),
                                                        ]
                                                    ),
                                                    #  ------------------------//----------------------//-------//
                                                ]
                                            )
                                        ),
                                        #  ----------------------for the transcriptions-----------------//
                                    ]
                                )
                            ),

                            #  -------------------the other container here-------------------//
                            ft.Container(
                                border_radius=ft.border_radius.all(10),
                                padding=ft.padding.only(left=30, right=10, top=10, bottom=30),
                                margin=ft.margin.only(right=20),
                                bgcolor="#f5f5f5",
                                content=ft.Row(
                                    alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                                    controls=[
                                        ft.Container(
                                            content=ft.Column(
                                                controls=[
                                                    ft.Text(
                                                        "recorded session".capitalize(),
                                                        style=ft.TextThemeStyle.DISPLAY_SMALL,
                                                        weight=ft.FontWeight.W_500,
                                                        color="#0078D9"
                                                    ),
                                                    ft.Row(
                                                        controls=[
                                                            ft.Container(
                                                                border_radius=ft.border_radius.all(30),
                                                                padding=ft.padding.only(left=50),
                                                                gradient=ft.LinearGradient(
                                                                    colors=[
                                                                        "#00B4C6",
                                                                        "#007BDD"
                                                                    ],
                                                                    begin=ft.alignment.bottom_left,
                                                                    end=ft.alignment.top_right
                                                                ),
                                                                width=500,
                                                                height=80,
                                                                content=ft.Row(
                                                                    controls=[
                                                                        ft.Icon(
                                                                            ft.icons.MUSIC_NOTE_ROUNDED,
                                                                            size=30,
                                                                            color="white",
                                                                        ),
                                                                        ft.Text(
                                                                            f"{self.recording_file_name.value}",
                                                                            size=18,
                                                                            weight=ft.FontWeight.W_700,
                                                                            color="white"
                                                                        ),
                                                                        ft.Container(
                                                                            content=ft.Row(
                                                                                controls=[
                                                                                    ft.Container(
                                                                                        content=ft.Row(
                                                                                            controls=[
                                                                                                ft.IconButton(
                                                                                                    icon=ft.icons.PLAY_ARROW_ROUNDED,
                                                                                                    icon_size=40,
                                                                                                    icon_color="black",
                                                                                                    tooltip="play",
                                                                                                    bgcolor="#f5f5f5",
                                                                                                    on_click=lambda
                                                                                                        _: self.audio1.play()
                                                                                                ),
                                                                                                ft.IconButton(
                                                                                                    icon=ft.icons.PAUSE_ROUNDED,
                                                                                                    icon_size=40,
                                                                                                    icon_color="black",
                                                                                                    bgcolor="#f5f5f5",
                                                                                                    tooltip="pause".title(),
                                                                                                    on_click=lambda
                                                                                                        _: self.audio1.pause()
                                                                                                ),
                                                                                                ft.IconButton(
                                                                                                    icon=ft.icons.VOLUME_DOWN_ROUNDED,
                                                                                                    icon_size=40,
                                                                                                    icon_color="black",
                                                                                                    bgcolor="#f5f5f5",
                                                                                                    tooltip="volume down".title(),
                                                                                                    on_click=self.volume_down
                                                                                                ),
                                                                                                ft.IconButton(
                                                                                                    icon=ft.icons.VOLUME_UP_ROUNDED,
                                                                                                    icon_size=40,
                                                                                                    icon_color="black",
                                                                                                    bgcolor="#f5f5f5",
                                                                                                    tooltip="volume up".title(),
                                                                                                    on_click=self.volume_up
                                                                                                ),
                                                                                            ]
                                                                                        )
                                                                                    )
                                                                                ]
                                                                            )
                                                                        ),

                                                                    ]
                                                                )
                                                            )
                                                        ]
                                                    ),

                                                ]
                                            )
                                        ),
                                        #  ---------------the container for the medea prayer and translations----//
                                        #  -----------------------the container for converting the text-----------//
                                        ft.Container(
                                            content=ft.Column(
                                                controls=[
                                                    ft.Text(
                                                        "translate".capitalize(),
                                                        style=ft.TextThemeStyle.DISPLAY_SMALL,
                                                        weight=ft.FontWeight.W_500,
                                                        color="#0078D9"
                                                    ),
                                                    ft.ElevatedButton(
                                                        height=80,
                                                        on_click=self.translate_voice,
                                                        bgcolor="#0078D9",
                                                        content=ft.Row(
                                                            controls=[
                                                                ft.Icon(ft.icons.G_TRANSLATE_ROUNDED, size=40,
                                                                        color="white"),
                                                                ft.Text(
                                                                    "translate generated voice".title(),
                                                                    color="white",
                                                                    size=20,
                                                                    weight=ft.FontWeight.W_500
                                                                )
                                                            ],

                                                        ),
                                                    )
                                                ]
                                            )
                                        )
                                    ]
                                )
                            ),
                        ]
                    )
                )
            ]
        )

    #  -----------------------//-----------------function to translate the voice to text here-------------------//
    def translate_voice(self, e):
        try:

            audio_file = open(self.recording_file_name.value, "rb")
            transcript = openai.Audio.transcribe("whisper-1", audio_file)
            translated_text = ft.AlertDialog(
                content=ft.Container(
                    content=ft.Column(
                        controls=[
                            ft.Text(
                                f"{transcript.text}"
                            )
                        ]
                    )
                )
            )
            self.page.dialog = translated_text
            translated_text.open = True
            self.page.update()
        except Exception as ex:
            print(ex)

    def fetch_all_recording(self):
        try:
            for recording_class in os.listdir(self.RECORDING_PATH):
                self.all_recordings.controls.append(
                    ft.Card(
                        elevation=0.5,
                        content=ft.Container(
                            bgcolor="#F2ECFF",
                            border_radius=ft.border_radius.all(10),
                            content=ft.Column(
                                controls=[
                                    #  ----------------the container for the circle avatar---------//
                                    ft.Container(
                                        margin=ft.margin.only(top=30),
                                        content=ft.Row(
                                            alignment=ft.MainAxisAlignment.CENTER,
                                            controls=[
                                                ft.CircleAvatar(
                                                    bgcolor="#007BDD",
                                                    width=150,
                                                    height=150,
                                                    content=ft.Row(
                                                        alignment=ft.MainAxisAlignment.CENTER,
                                                        controls=[
                                                            ft.Icon(
                                                                ft.icons.MUSIC_NOTE_ROUNDED,
                                                                size=50,
                                                                color="#FFB84A"
                                                            )
                                                        ]
                                                    )
                                                ),

                                            ]
                                        )
                                    ),
                                    ft.Container(
                                        content=ft.Row(
                                            alignment=ft.MainAxisAlignment.CENTER,
                                            controls=[
                                                ft.Text(f"{recording_class}")
                                            ]
                                        )
                                    ),
                                    ft.Container(
                                        margin=ft.margin.only(top=10),
                                        padding=ft.padding.only(left=20),
                                        content=ft.Row(
                                            controls=[
                                                ft.ElevatedButton(
                                                    data=recording_class,
                                                    on_click=self.on_click,
                                                    bgcolor="#007BDD",
                                                    content=ft.Row(
                                                        controls=[
                                                            ft.Icon(
                                                                ft.icons.BROADCAST_ON_HOME_ROUNDED,
                                                                size=20,
                                                                color="white"),
                                                            ft.Text(
                                                                "translate recording".title(),
                                                                color="white"
                                                            )
                                                        ],
                                                    ),
                                                ),
                                            ]
                                        )
                                    )
                                ]
                            )
                        )
                    )
                )

        except Exception as ex:
            print(ex)

    #  -----------------------the onclick control page----------------------//
    def on_click(self, e):
        try:
            self.page.snack_bar = ft.SnackBar(
                bgcolor="#0050C1",
                content=ft.Container(
                    content=ft.Row(
                        controls=[
                            ft.Text(
                                "your translation has started and you will see the output".capitalize(),
                                style=ft.TextThemeStyle.DISPLAY_SMALL,
                                weight=ft.FontWeight.W_700,
                                color="white"
                            )
                        ]
                    )
                )
            )
            self.page.snack_bar.open = True
            self.page.update()
            self.recording_name = e.control.data
            audio_file = open(f"C:/FinalProject/recordings/{self.recording_name}", "rb")
            transcript = openai.Audio.transcribe("whisper-1", audio_file)
            results = transcript['text']
            #  -------------------printing the translated text to word document----------//
            self.translated_text = ft.AlertDialog(
                content=ft.Container(
                    border_radius=ft.border_radius.all(10),
                    width=600,
                    bgcolor="white",
                    padding=ft.padding.only(left=20, top=20),
                    content=ft.Column(
                        controls=[
                            ft.Text(
                                "generated translation".title(),
                                weight=ft.FontWeight.BOLD,
                                color="#0050C1",
                                style=ft.TextThemeStyle.DISPLAY_MEDIUM

                            ),
                            ft.Container(
                                margin=ft.margin.only(top=10),
                                content=ft.Column(
                                    controls=[
                                        ft.Text(
                                            f"{results}",
                                            style=ft.TextThemeStyle.DISPLAY_SMALL,
                                            weight=ft.FontWeight.BOLD
                                        )
                                    ]
                                )
                            ),
                            ft.Container(
                              content=ft.Row(
                                  controls=[
                                      self.document_name
                                  ]
                              )
                            ),
                            ft.Container(
                                content=ft.Row(
                                    controls=[
                                        ft.ElevatedButton(
                                            on_click=self.export_to_word,
                                            bgcolor="#007BDD",
                                            content=ft.Row(
                                                controls=[
                                                    ft.Icon(
                                                        ft.icons.BROADCAST_ON_HOME_ROUNDED,
                                                        size=20,
                                                        color="white"),
                                                    ft.Text(
                                                        "save to word".title(),
                                                        color="white"
                                                    )
                                                ],
                                            ),
                                        ),
                                    ]
                                )
                            )
                        ]
                    )
                ),
            )
            self.page.dialog = self.translated_text
            self.translated_text.open = True
            self.page.update()
        except Exception as ex:
            print(ex)

    # - ------------------------exporting to word----------------------//
    def export_to_word(self, e):
        try:
            audio_file = open(f"C:/FinalProject/recordings/{self.recording_name}", "rb")
            transcript = openai.Audio.transcribe("whisper-1", audio_file)
            results = transcript['text']

            document = Document()
            # add a heading to the document
            document.add_heading('SOCHIE TECHNICAL COLLEGE', 0)
            # add a paragraph to the document
            p1 = document.add_paragraph('THE LECTURE LESSON')
            # add another paragraph to the document
            p2 = document.add_paragraph('This is another paragraph.')

            p3 = document.add_paragraph(f"{results}")
            # add an image to the document
            img_path = 'assets/signs/pexels-andre-moura-2402467.jpg'
            document.add_picture(img_path, width=Inches(2))
            # save the document
            document.save(f'C:/FinalProject/Lessons/{self.document_name.value}.docx')

            self.translated_text.open = False
            self.page.update()
            self.page.snack_bar = ft.SnackBar(
                bgcolor="#0050C1",
                content=ft.Container(
                    content=ft.Row(
                        controls=[
                            ft.Text(
                                "your document has been saved successfully".capitalize(),
                                style=ft.TextThemeStyle.DISPLAY_SMALL,
                                weight=ft.FontWeight.W_700,
                                color="white"
                            )
                        ]
                    )
                )
            )
            self.page.snack_bar.open = True
            self.page.update()
            self.fetch_all_recording()
        except Exception as ex:
            print(ex)
